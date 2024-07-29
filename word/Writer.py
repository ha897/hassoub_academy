import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
"""
كتابة قصة
"""
file = docx.Document()
#عنوان
textH = file.add_heading('أحمد: الطفل الذي أصبح كاتبًا')
textH.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#نص
text = "في قرية صغيرة، كان هناك طفل يدعى أحمد. كان يحب القراءة كثيراً وكان يحلم بأن يصبح كاتباً يوماً ما. في أحد الأيام، وجد كتاب قديم في مكتبة جده. كان الكتاب مليء بالقصص المدهشة عن الأماكن البعيدة والمغامرات الشيقة. قرأ أحمد الكتاب من الغلاف إلى الغلاف، وكل قصة جعلته يحلم أكثر وأكثر بأن يصبح كاتباً. بعد سنوات من العمل الشاق والكتابة، أصبح أحمد أحد أشهر الكتاب في بلده. وقد استوحى كل قصصه من الكتاب القديم الذي وجده في طفولته."
textA = file.add_paragraph(text)
textA.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#صورة
file.add_picture('pic.jpeg')

#عنوان
textHE = file.add_heading('Ahmed: The Boy Who Became a Writer')
textHE.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#نص
text = "In a small village, there was a boy named Ahmed. He loved reading very much and dreamed of becoming a writer one day. One day, he found an old book in his grandfather’s library. The book was filled with amazing stories about distant places and exciting adventures. Ahmed read the book from cover to cover, and each story made him dream more and more of becoming a writer. After years of hard work and writing, Ahmed became one of the most famous writers in his country. He drew all his stories from the old book he found in his childhood."
file.add_paragraph(text)
#صورة
file.add_picture('pic.jpeg')


file.save('new2.docx')